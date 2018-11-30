# -*- coding: utf-8 -*-
import os

# Imports the Google Cloud client library
from google.cloud import translate
from docx import Document

print('Credentials from environ: {}'.format(os.environ.get('GOOGLE_APPLICATION_CREDENTIALS')))


# Instantiates a client
translate_client = translate.Client()


def translate_paragraph(text, target_language='en'):
    translation = translate_client.translate(
        p.text,
        target_language='en')
    return translation['translatedText']


if __name__ == "__main__":

    in_filepath = 'in/Annexe 2 CGA FFF  2017.docx'
    out_filepath = 'out/output.docx'

    doc = Document(in_filepath)
    count = 1
    total = len(doc.paragraphs)

    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    total += 1

    for p in doc.paragraphs:
        p.text = translate_paragraph(p.text)
        print 'Paragraph %s of %s: %s' % (count, total, p.text)
        count += 1

    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    p.text = translate_paragraph(p.text)
                    print 'Paragraph %s of %s: %s' % (count, total, p.text)
                    count += 1

    doc.save(out_filepath)